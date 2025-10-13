"""Lesson plan export utilities - exports lesson plans to document formats"""

import datetime
import tempfile
import os
import json
from typing import Dict, List, Tuple, Any, Optional, Union


class LessonExporter:
    """Export lesson plans to various document formats"""
    
    @staticmethod
    def smart_export(lesson_plans: List[Union[str, Dict]], 
                     course_outline: Optional[Dict] = None,
                     template_mode: str = "text",
                     template_path: Optional[str] = None,
                     export_format: str = "word") -> Tuple[Optional[str], bool]:
        """
        智能导出：自动选择填充模板或生成新文档
        
        Args:
            lesson_plans: 教案列表（可以是字符串或字典）
            course_outline: 课程大纲信息
            template_mode: 模板类型 ("tags" 或 "text")
            template_path: 原始模板路径
            export_format: 导出格式 ("word" 或 "txt")
            
        Returns:
            Tuple of (file_path, success)
        """
        try:
            print(f"📤 智能导出开始...")
            print(f"   - 模板模式: {template_mode}")
            print(f"   - 教案数量: {len(lesson_plans)}")
            print(f"   - 导出格式: {export_format}")
            
            # 判断是否为标签模式且有模板
            if template_mode == "tags" and template_path and all(isinstance(lp, dict) for lp in lesson_plans):
                print(f"✅ 使用标签填充模式")
                return LessonExporter.export_with_template_filling(
                    lesson_plans, template_path, export_format
                )
            else:
                print(f"✅ 使用传统生成模式")
                # 转换为字符串列表（如果是字典，转为文本）
                text_plans = []
                for lp in lesson_plans:
                    if isinstance(lp, dict):
                        # 将字典转为易读文本
                        text_plans.append(LessonExporter._dict_to_text(lp))
                    else:
                        text_plans.append(str(lp))
                
                if export_format == "word":
                    return LessonExporter.export_to_word(text_plans, course_outline)
                else:
                    return LessonExporter.export_to_txt(text_plans, course_outline)
                    
        except Exception as e:
            print(f"❌ 智能导出失败: {e}")
            import traceback
            traceback.print_exc()
            return None, False
    
    @staticmethod
    def export_with_template_filling(lesson_plans: List[Dict], 
                                     template_path: str,
                                     export_format: str = "word") -> Tuple[Optional[str], bool]:
        """
        使用模板填充方式导出（针对标签模式）
        
        Args:
            lesson_plans: 教案数据字典列表
            template_path: 模板文件路径
            export_format: 导出格式
            
        Returns:
            Tuple of (file_path, success)
        """
        try:
            from utils.template_filler import WordTemplateFiller
            
            print(f"🏷️  开始使用模板填充导出...")
            filler = WordTemplateFiller()
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            exports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'exports')
            os.makedirs(exports_dir, exist_ok=True)
            
            # 如果只有一个教案，直接填充
            if len(lesson_plans) == 1:
                filename = f"lesson_plan_{timestamp}.docx"
                output_path = os.path.join(exports_dir, filename)
                
                success = filler.fill_template(
                    template_path=template_path,
                    output_path=output_path,
                    data=lesson_plans[0]
                )
                
                if success:
                    print(f"✅ 单个教案填充成功: {output_path}")
                    return output_path, True
                else:
                    print(f"❌ 教案填充失败")
                    return None, False
            
            # 多个教案：合并或分别生成
            else:
                # 方案1：合并成一个文档（按课次）
                from docx import Document
                from docxtpl import DocxTemplate
                import shutil
                
                filename = f"all_lesson_plans_{timestamp}.docx"
                output_path = os.path.join(exports_dir, filename)
                
                # 复制第一个教案作为基础
                first_output = os.path.join(exports_dir, f"temp_lesson_1.docx")
                filler.fill_template(template_path, first_output, lesson_plans[0])
                
                # 加载第一个文档
                combined_doc = Document(first_output)
                
                # 为后续教案添加分页和内容
                for i, lesson_data in enumerate(lesson_plans[1:], 2):
                    # 生成单个教案
                    temp_output = os.path.join(exports_dir, f"temp_lesson_{i}.docx")
                    filler.fill_template(template_path, temp_output, lesson_data)
                    
                    # 添加分页
                    combined_doc.add_page_break()
                    
                    # 读取新教案内容并添加
                    lesson_doc = Document(temp_output)
                    for element in lesson_doc.element.body:
                        combined_doc.element.body.append(element)
                    
                    # 删除临时文件
                    os.remove(temp_output)
                
                # 保存合并文档
                combined_doc.save(output_path)
                
                # 删除第一个临时文件
                os.remove(first_output)
                
                print(f"✅ {len(lesson_plans)} 个教案合并填充成功: {output_path}")
                return output_path, True
                
        except Exception as e:
            print(f"❌ 模板填充导出失败: {e}")
            import traceback
            traceback.print_exc()
            return None, False
    
    @staticmethod
    def _dict_to_text(data: Dict) -> str:
        """将字典转换为易读文本"""
        lines = []
        for key, value in data.items():
            # 格式化键名
            formatted_key = key.replace('_', ' ').title()
            lines.append(f"**{formatted_key}**: {value}\n")
        return "\n".join(lines)
    
    @staticmethod
    def export_to_word(lesson_plans: List[str], course_outline: Optional[Dict] = None) -> Tuple[Optional[str], bool]:
        """
        Export lesson plans to Word document
        
        Args:
            lesson_plans: List of lesson plan contents
            course_outline: Optional course outline information
            
        Returns:
            Tuple of (file_path, success)
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"university_lesson_plans_{timestamp}.docx"
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('大学教案生成系统 - 教案导出', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(10)
            meta_para.add_run(f"\n共计教案：{len(lesson_plans)} 个").font.size = Pt(10)
            
            doc.add_paragraph()
            
            # Add course information
            if course_outline:
                course_info = course_outline.get('course_info', {})
                
                doc.add_heading('课程基本信息', 1)
                info_table = doc.add_table(rows=4, cols=2)
                info_table.style = 'Light Grid Accent 1'
                
                info_table.rows[0].cells[0].text = '课程名称'
                info_table.rows[0].cells[1].text = course_info.get('course_name', '')
                info_table.rows[1].cells[0].text = '课程性质'
                info_table.rows[1].cells[1].text = course_info.get('course_type', '')
                info_table.rows[2].cells[0].text = '学分学时'
                info_table.rows[2].cells[1].text = f"{course_info.get('credits', '')}学分 / {course_info.get('total_hours', '')}学时"
                info_table.rows[3].cells[0].text = '授课对象'
                info_table.rows[3].cells[1].text = course_info.get('target_students', '')
                
                doc.add_paragraph()
            
            # Add each lesson plan
            for i, plan in enumerate(lesson_plans):
                if i > 0:
                    doc.add_page_break()
                
                doc.add_heading(f'第 {i+1} 次课教案', 1)
                LessonExporter._add_content_to_doc(doc, plan)
            
            # Save to exports directory
            exports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'exports')
            os.makedirs(exports_dir, exist_ok=True)
            file_path = os.path.join(exports_dir, filename)
            doc.save(file_path)
            
            print(f"✅ Word文档导出成功: {file_path}")
            print(f"📁 文件保存位置: {os.path.abspath(file_path)}")
            return file_path, True
            
        except ImportError as e:
            print(f"⚠️ python-docx未安装，使用TXT格式: {e}")
            # Fallback to txt if python-docx not installed
            return LessonExporter.export_to_txt(lesson_plans, course_outline)
        except Exception as e:
            print(f"❌ Word导出失败: {e}")
            import traceback
            traceback.print_exc()
            return None, False
    
    @staticmethod
    def export_to_txt(lesson_plans: List[str], course_outline: Optional[Dict] = None) -> Tuple[Optional[str], bool]:
        """
        Export lesson plans to plain text file
        
        Args:
            lesson_plans: List of lesson plan contents
            course_outline: Optional course outline information
            
        Returns:
            Tuple of (file_path, success)
        """
        try:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"university_lesson_plans_{timestamp}.txt"
            
            content = f"{'='*80}\n"
            content += f"大学教案生成系统 - 教案导出\n"
            content += f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            content += f"共计教案：{len(lesson_plans)} 个\n"
            content += f"{'='*80}\n\n"
            
            if course_outline:
                course_info = course_outline.get('course_info', {})
                content += f"课程信息：\n"
                content += f"  课程名称：{course_info.get('course_name', '')}\n"
                content += f"  课程性质：{course_info.get('course_type', '')}\n"
                content += f"  学分学时：{course_info.get('credits', '')}学分 / {course_info.get('total_hours', '')}学时\n"
                content += f"  授课对象：{course_info.get('target_students', '')}\n\n"
            
            for i, plan in enumerate(lesson_plans):
                content += f"{'='*60}\n"
                content += f"第 {i+1} 次课教案\n"
                content += f"{'='*60}\n\n"
                
                # Clean up plan content
                cleaned_plan = LessonExporter._clean_content(plan)
                content += cleaned_plan + "\n\n"
            
            # Save to exports directory
            exports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'exports')
            os.makedirs(exports_dir, exist_ok=True)
            file_path = os.path.join(exports_dir, filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            print(f"✅ TXT文档导出成功: {file_path}")
            print(f"📁 文件保存位置: {os.path.abspath(file_path)}")
            return file_path, True
        except Exception as e:
            print(f"❌ TXT导出失败: {e}")
            return None, False
    
    @staticmethod
    def _add_content_to_doc(doc, content: str):
        """Add lesson plan content to Word document with formatting"""
        from utils.json_parser import fix_and_extract_json
        
        # Try to parse JSON if present
        parsed_content = content
        try:
            if content.strip().startswith('{') or content.strip().startswith('['):
                # Try to extract JSON
                json_data = fix_and_extract_json(content)
                if json_data:
                    parsed_content = LessonExporter._json_to_markdown(json_data)
                else:
                    parsed_content = content
        except Exception as e:
            print(f"JSON parsing skipped: {e}")
            parsed_content = content
        
        # Process markdown-style content
        lines = parsed_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)
            # Bullet lists
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered lists
            elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                doc.add_paragraph(line[3:], style='List Number')
            # Bold text
            elif '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
            # Normal paragraph
            else:
                doc.add_paragraph(line)
    
    @staticmethod
    def _json_to_markdown(data: Any, indent: int = 0) -> str:
        """Convert JSON structure to markdown text"""
        text_parts = []
        indent_str = '  ' * indent
        
        if isinstance(data, dict):
            for key, value in data.items():
                # Format key name
                formatted_key = key.replace('_', ' ').title()
                
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{indent_str}## {formatted_key}")
                    text_parts.append(LessonExporter._json_to_markdown(value, indent + 1))
                else:
                    text_parts.append(f"{indent_str}**{formatted_key}**: {value}")
        
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, (dict, list)):
                    text_parts.append(LessonExporter._json_to_markdown(item, indent))
                else:
                    text_parts.append(f"{indent_str}- {item}")
        
        else:
            text_parts.append(f"{indent_str}{data}")
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def _clean_content(content: str) -> str:
        """Clean up content for plain text export"""
        # Remove excessive markdown formatting
        cleaned = content.replace('**', '')
        cleaned = cleaned.replace('##', '')
        cleaned = cleaned.replace('#', '')
        
        # Try to parse JSON and convert to readable text
        try:
            if cleaned.strip().startswith('{'):
                from utils.json_parser import fix_and_extract_json
                json_data = fix_and_extract_json(cleaned)
                if json_data:
                    cleaned = LessonExporter._json_to_text(json_data)
        except:
            pass
        
        return cleaned
    
    @staticmethod
    def _json_to_text(data: Any, indent: int = 0) -> str:
        """Convert JSON to plain readable text"""
        text_parts = []
        indent_str = '  ' * indent
        
        if isinstance(data, dict):
            for key, value in data.items():
                formatted_key = key.replace('_', ' ').title()
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{indent_str}{formatted_key}:")
                    text_parts.append(LessonExporter._json_to_text(value, indent + 1))
                else:
                    text_parts.append(f"{indent_str}{formatted_key}: {value}")
        
        elif isinstance(data, list):
            for i, item in enumerate(data, 1):
                if isinstance(item, (dict, list)):
                    text_parts.append(LessonExporter._json_to_text(item, indent))
                else:
                    text_parts.append(f"{indent_str}{i}. {item}")
        
        else:
            text_parts.append(f"{indent_str}{data}")
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def get_export_formats() -> List[str]:
        """Get available export formats"""
        formats = ['txt']
        try:
            import docx
            formats.insert(0, 'docx')
        except ImportError:
            pass
        return formats