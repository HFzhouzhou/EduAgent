"""
Word 模板填充工具 - 高级功能
Template Filler - Advanced Feature

支持通过 XML 标签将内容填充到 Word 模板中
使用 docxtpl 库实现 Jinja2 模板语法
"""

from pathlib import Path
from typing import Dict, Any, Optional, List
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json


class WordTemplateFiller:
    """Word 模板填充器"""
    
    def __init__(self):
        self.supported_tags = [
            # 基本信息
            'course_name',              # 课程名称
            'teacher_name',             # 授课教师
            'class_name',               # 授课班级
            'lesson_number',            # 课次
            'lesson_title',             # 课题
            'teaching_hours',           # 学时
            'chapter_section',          # 授课章节
            
            # 教学目标
            'ideological_goals',        # 思政育人目标
            'knowledge_goals',          # 知识目标
            'ability_goals',            # 能力目标
            'ideological_elements',     # 思政元素
            
            # 教学重难点
            'teaching_focus',           # 教学重点
            'focus_solutions',          # 教学重点解决措施
            'teaching_difficulty',      # 教学难点
            'difficulty_solutions',     # 教学难点解决措施
            
            # 教学方法
            'teaching_methods',         # 教法
            'learning_methods',         # 学法
            'teaching_resources',       # 教学资源
            
            # 课前预习
            'preview_content',          # 课前预习-教学内容
            'preview_teacher',          # 课前预习-教师活动
            'preview_student',          # 课前预习-学生活动
            'preview_intention',        # 课前预习-设计意图
            
            # 自主学习
            'self_learning_content',    # 自主学习-预习内容
            'self_learning_teacher',    # 自主学习-教师活动
            'self_learning_student',    # 自主学习-学生活动
            'self_learning_intention',  # 自主学习-设计意图
            
            # 新课导入
            'introduction_content',     # 新课导入-教学内容
            'introduction_teacher',     # 新课导入-教师活动
            'introduction_student',     # 新课导入-学生活动
            'introduction_intention',   # 新课导入-设计意图
            
            # 预习反馈
            'feedback_content',         # 预习反馈-教学内容
            'feedback_teacher',         # 预习反馈-教师活动
            'feedback_student',         # 预习反馈-学生活动
            'feedback_intention',       # 预习反馈-设计意图
            
            # 新课讲授
            'teaching_content',         # 新课讲授-教学内容
            'teaching_teacher',         # 新课讲授-教师活动
            'teaching_student',         # 新课讲授-学生活动
            'teaching_intention',       # 新课讲授-设计意图
            
            # 实践环节
            'practice_content',         # 实践-教学内容
            'practice_teacher',         # 实践-教师活动
            'practice_student',         # 实践-学生活动
            'practice_intention',       # 实践-设计意图
            
            # 展示环节
            'presentation_content',     # 展示-教学内容
            'presentation_teacher',     # 展示-教师活动
            'presentation_student',     # 展示-学生活动
            'presentation_intention',   # 展示-设计意图
            
            # 评价环节
            'evaluation_content',       # 评价-教学内容
            'evaluation_teacher',       # 评价-教师活动
            'evaluation_student',       # 评价-学生活动
            'evaluation_intention',     # 评价-设计意图
            
            # 课后作业
            'homework_content',         # 课后作业-教学内容
            'homework_teacher',         # 课后作业-教师活动
            'homework_student',         # 课后作业-学生活动
            'homework_intention',       # 课后作业-设计意图
            
            # 阅读延伸
            'extension_content',        # 阅读延伸-教学内容
            'extension_teacher',        # 阅读延伸-教师活动
            'extension_student',        # 阅读延伸-学生活动
            'extension_intention',      # 阅读延伸-设计意图
            
            # 教学反思
            'reflection_effects',       # 教学反思-目标效果
            'reflection_improvements',  # 教学反思-反思改进
        ]
    
    def get_tag_guide(self) -> str:
        """获取标签使用指南"""
        guide = """
# Word 模板标签使用指南（完整版）

## 📝 支持的标签列表（共72个）

### 一、基本信息（7个）
- `{{course_name}}` - 课程名称
- `{{teacher_name}}` - 授课教师
- `{{class_name}}` - 授课班级
- `{{lesson_number}}` - 课次
- `{{lesson_title}}` - 课题
- `{{teaching_hours}}` - 学时
- `{{chapter_section}}` - 授课章节

### 二、教学目标（4个）
- `{{ideological_goals}}` - 思政育人目标
- `{{knowledge_goals}}` - 知识目标
- `{{ability_goals}}` - 能力目标
- `{{ideological_elements}}` - 思政元素

### 三、教学重难点（4个）
- `{{teaching_focus}}` - 教学重点
- `{{focus_solutions}}` - 教学重点解决措施
- `{{teaching_difficulty}}` - 教学难点
- `{{difficulty_solutions}}` - 教学难点解决措施

### 四、教学方法（3个）
- `{{teaching_methods}}` - 教法
- `{{learning_methods}}` - 学法
- `{{teaching_resources}}` - 教学资源

### 五、课前预习（4个）
- `{{preview_content}}` - 教学内容
- `{{preview_teacher}}` - 教师活动
- `{{preview_student}}` - 学生活动
- `{{preview_intention}}` - 设计意图

### 六、自主学习（4个）
- `{{self_learning_content}}` - 预习内容
- `{{self_learning_teacher}}` - 教师活动
- `{{self_learning_student}}` - 学生活动
- `{{self_learning_intention}}` - 设计意图

### 七、新课导入（4个）
- `{{introduction_content}}` - 教学内容
- `{{introduction_teacher}}` - 教师活动
- `{{introduction_student}}` - 学生活动
- `{{introduction_intention}}` - 设计意图

### 八、预习反馈（4个）
- `{{feedback_content}}` - 教学内容
- `{{feedback_teacher}}` - 教师活动
- `{{feedback_student}}` - 学生活动
- `{{feedback_intention}}` - 设计意图

### 九、新课讲授（4个）
- `{{teaching_content}}` - 教学内容
- `{{teaching_teacher}}` - 教师活动
- `{{teaching_student}}` - 学生活动
- `{{teaching_intention}}` - 设计意图

### 十、实践环节（4个）
- `{{practice_content}}` - 教学内容
- `{{practice_teacher}}` - 教师活动
- `{{practice_student}}` - 学生活动
- `{{practice_intention}}` - 设计意图

### 十一、展示环节（4个）
- `{{presentation_content}}` - 教学内容
- `{{presentation_teacher}}` - 教师活动
- `{{presentation_student}}` - 学生活动
- `{{presentation_intention}}` - 设计意图

### 十二、评价环节（4个）
- `{{evaluation_content}}` - 教学内容
- `{{evaluation_teacher}}` - 教师活动
- `{{evaluation_student}}` - 学生活动
- `{{evaluation_intention}}` - 设计意图

### 十三、课后作业（4个）
- `{{homework_content}}` - 教学内容
- `{{homework_teacher}}` - 教师活动
- `{{homework_student}}` - 学生活动
- `{{homework_intention}}` - 设计意图

### 十四、阅读延伸（4个）
- `{{extension_content}}` - 教学内容
- `{{extension_teacher}}` - 教师活动
- `{{extension_student}}` - 学生活动
- `{{extension_intention}}` - 设计意图

### 十五、教学反思（2个）
- `{{reflection_effects}}` - 目标效果
- `{{reflection_improvements}}` - 反思改进

## 🎯 使用方法

### 步骤 1: 在 Word 中插入标签

1. 打开您的教案模板（.docx）
2. 在需要填充内容的位置，输入标签，例如：`{{course_name}}`
3. 保存模板

### 步骤 2: 上传模板到系统

在系统中上传这个带标签的模板

### 步骤 3: 生成教案

系统会自动识别标签并填充对应内容

## 💡 示例

### 模板示例：

```
课程名称：{{course_name}}
授课教师：{{teacher_name}}
授课班级：{{class_name}}

第 {{lesson_number}} 次课

一、课题
{{lesson_title}}

二、教学目标
{{teaching_objectives}}

三、教学重点
{{teaching_focus}}

四、教学难点
{{teaching_difficulty}}

五、教学过程
{{teaching_process}}

六、课后作业
{{homework}}
```

### 填充后效果：

```
课程名称：数据结构
授课教师：张三
授课班级：计算机2021级1班

第 1 次课

一、课题
数据结构概述

二、教学目标
1. 掌握数据结构的基本概念
2. 理解数据结构的分类
...
```

## ⚠️ 注意事项

1. 标签必须使用 `{{}}` 包裹
2. 标签名必须与支持列表中的一致
3. 标签区分大小写
4. 可以在任何位置使用标签（表格、文本框等）
5. 同一个标签可以多次使用

## 🔧 高级用法

### 条件判断
```
{% if homework %}
课后作业：{{homework}}
{% endif %}
```

### 循环列表
```
教学目标：
{% for objective in objectives %}
- {{objective}}
{% endfor %}
```

### 格式化
```
授课日期：{{date | date_format}}
```

---

完整文档请访问：https://docxtpl.readthedocs.io/
"""
        return guide
    
    def fill_template(
        self,
        template_path: str,
        output_path: str,
        data: Dict[str, Any]
    ) -> bool:
        """
        填充 Word 模板
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data: 要填充的数据字典
            
        Returns:
            bool: 是否成功
        """
        try:
            # 加载模板
            doc = DocxTemplate(template_path)
            
            # 填充数据
            doc.render(data)
            
            # 保存文件
            doc.save(output_path)
            
            print(f"✅ 模板填充成功：{output_path}")
            return True
            
        except Exception as e:
            print(f"❌ 模板填充失败：{e}")
            return False
    
    def fill_lesson_plan(
        self,
        template_path: str,
        output_path: str,
        lesson_plan: Dict[str, Any],
        course_info: Dict[str, Any]
    ) -> bool:
        """
        填充教案到模板
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            lesson_plan: 教案数据
            course_info: 课程信息
            
        Returns:
            bool: 是否成功
        """
        # 准备数据
        data = {
            # 课程基本信息
            'course_name': course_info.get('course_name', ''),
            'teacher_name': course_info.get('teacher_name', ''),
            'class_name': course_info.get('class_name', ''),
            
            # 课次信息
            'lesson_number': lesson_plan.get('lesson_number', ''),
            'lesson_title': lesson_plan.get('lesson_title', ''),
            'teaching_hours': lesson_plan.get('teaching_hours', ''),
            
            # 教学内容
            'teaching_objectives': self._format_objectives(
                lesson_plan.get('teaching_objectives', {})
            ),
            'teaching_focus': self._format_list(
                lesson_plan.get('teaching_focus', [])
            ),
            'teaching_difficulty': self._format_list(
                lesson_plan.get('teaching_difficulty', [])
            ),
            'teaching_methods': self._format_list(
                lesson_plan.get('teaching_methods', [])
            ),
            
            # 教学过程
            'teaching_process': self._format_teaching_process(
                lesson_plan.get('teaching_process', {})
            ),
            
            # 其他
            'homework': lesson_plan.get('homework', ''),
            'reflection': lesson_plan.get('reflection', ''),
        }
        
        return self.fill_template(template_path, output_path, data)
    
    def _format_objectives(self, objectives: Dict[str, List[str]]) -> str:
        """格式化教学目标"""
        if not objectives:
            return ""
        
        result = []
        for category, items in objectives.items():
            if items:
                result.append(f"\n{category}：")
                for i, item in enumerate(items, 1):
                    result.append(f"{i}. {item}")
        
        return "\n".join(result)
    
    def _format_list(self, items: List[str]) -> str:
        """格式化列表"""
        if not items:
            return ""
        
        return "\n".join(f"{i}. {item}" for i, item in enumerate(items, 1))
    
    def _format_teaching_process(self, process: Dict[str, Any]) -> str:
        """格式化教学过程"""
        if not process:
            return ""
        
        result = []
        for step_name, step_content in process.items():
            result.append(f"\n【{step_name}】")
            if isinstance(step_content, list):
                for item in step_content:
                    result.append(f"• {item}")
            else:
                result.append(str(step_content))
        
        return "\n".join(result)
    
    def check_template_tags(self, template_path: str) -> Dict[str, Any]:
        """
        检查模板中的标签
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            dict: 包含找到的标签和未识别的标签
        """
        try:
            doc = DocxTemplate(template_path)
            
            # 获取模板中的所有变量
            undeclared_vars = doc.undeclared_template_variables
            
            # 分类标签
            recognized_tags = []
            unrecognized_tags = []
            
            for var in undeclared_vars:
                if var in self.supported_tags:
                    recognized_tags.append(var)
                else:
                    unrecognized_tags.append(var)
            
            return {
                'success': True,
                'has_tags': len(undeclared_vars) > 0,
                'recognized_tags': recognized_tags,
                'unrecognized_tags': unrecognized_tags,
                'total_tags': len(undeclared_vars)
            }
            
        except Exception as e:
            return {
                'success': False,
                'has_tags': False,
                'error': str(e)
            }
    
    def detect_template_mode(self, template_path: str) -> str:
        """
        快速检测模板模式
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            str: "tags" 或 "text"
        """
        try:
            result = self.check_template_tags(template_path)
            if result.get('success') and result.get('has_tags'):
                return "tags"
            return "text"
        except:
            return "text"
    
    def create_sample_template(self, output_path: str) -> bool:
        """
        创建示例模板
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            bool: 是否成功
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 标题
            title = doc.add_heading('教案模板（带智能标签）', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 课程信息表格
            table = doc.add_table(rows=3, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # 第一行
            table.rows[0].cells[0].text = '课程名称'
            table.rows[0].cells[1].text = '{{course_name}}'
            table.rows[0].cells[2].text = '授课教师'
            table.rows[0].cells[3].text = '{{teacher_name}}'
            
            # 第二行
            table.rows[1].cells[0].text = '授课班级'
            table.rows[1].cells[1].text = '{{class_name}}'
            table.rows[1].cells[2].text = '课次'
            table.rows[1].cells[3].text = '第 {{lesson_number}} 次课'
            
            # 第三行
            table.rows[2].cells[0].text = '课题'
            table.rows[2].cells[1].text = '{{lesson_title}}'
            table.rows[2].cells[2].text = '学时'
            table.rows[2].cells[3].text = '{{teaching_hours}}'
            
            # 教学目标
            doc.add_heading('一、教学目标', 1)
            doc.add_paragraph('{{teaching_objectives}}')
            
            # 教学重点
            doc.add_heading('二、教学重点', 1)
            doc.add_paragraph('{{teaching_focus}}')
            
            # 教学难点
            doc.add_heading('三、教学难点', 1)
            doc.add_paragraph('{{teaching_difficulty}}')
            
            # 教学方法
            doc.add_heading('四、教学方法', 1)
            doc.add_paragraph('{{teaching_methods}}')
            
            # 教学过程
            doc.add_heading('五、教学过程', 1)
            doc.add_paragraph('{{teaching_process}}')
            
            # 课后作业
            doc.add_heading('六、课后作业', 1)
            doc.add_paragraph('{{homework}}')
            
            # 教学反思
            doc.add_heading('七、教学反思', 1)
            doc.add_paragraph('{{reflection}}')
            
            # 保存
            doc.save(output_path)
            
            print(f"✅ 示例模板创建成功：{output_path}")
            return True
            
        except Exception as e:
            print(f"❌ 创建示例模板失败：{e}")
            return False


# 使用示例
if __name__ == "__main__":
    filler = WordTemplateFiller()
    
    # 创建示例模板
    sample_path = "教案模板_智能标签版.docx"
    filler.create_sample_template(sample_path)
    
    # 打印使用指南
    print(filler.get_tag_guide())
    
    # 示例：填充教案
    lesson_data = {
        'lesson_number': '1',
        'lesson_title': '数据结构概述',
        'teaching_hours': '2',
        'teaching_objectives': {
            '知识目标': [
                '掌握数据结构的基本概念',
                '理解数据结构的分类',
            ],
            '能力目标': [
                '能够分析问题的数据结构特点',
            ]
        },
        'teaching_focus': [
            '数据结构的定义',
            '数据结构的分类',
        ],
        'teaching_difficulty': [
            '数据结构的抽象表示',
        ],
        'teaching_methods': [
            '讲授法',
            '案例分析法',
        ],
        'teaching_process': {
            '导入环节': ['回顾程序设计基础', '引出数据结构概念'],
            '讲授环节': ['介绍数据结构定义', '讲解分类方法'],
            '练习环节': ['分析实际案例', '小组讨论'],
        },
        'homework': '阅读教材第一章，完成课后习题1-5',
        'reflection': '（课后填写）',
    }
    
    course_data = {
        'course_name': '数据结构',
        'teacher_name': '张老师',
        'class_name': '计算机2021级1班',
    }
    
    # filler.fill_lesson_plan(
    #     template_path=sample_path,
    #     output_path="教案_第1次课.docx",
    #     lesson_plan=lesson_data,
    #     course_info=course_data
    # )

